from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
import os
import pdfplumber
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from docx.shared import Inches
import pytesseract

# Set Tesseract path (for OCR)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

app = Flask(__name__)

UPLOAD_DIR = 'static/uploads'
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.route('/', methods=['GET', 'POST'])
def home():
    if request.method == 'POST':
        uploaded_file = request.files.get('pdf')
        if uploaded_file and uploaded_file.filename.lower().endswith('.pdf'):
            saved_path = os.path.join(UPLOAD_DIR, secure_filename(uploaded_file.filename))
            uploaded_file.save(saved_path)

            output_docx = process_pdf_to_word(saved_path)
            return send_file(output_docx, as_attachment=True)
    return render_template('index.html')

def process_pdf_to_word(pdf_file_path):
    word_doc = Document()

    # Extract text (if available) from each PDF page
    with pdfplumber.open(pdf_file_path) as pdf_pages:
        for page in pdf_pages.pages:
            page_text = page.extract_text()
            if page_text:
                word_doc.add_paragraph(page_text)

    # Convert pages to images for OCR and visual embedding
    page_images = convert_from_path(pdf_file_path)
    for idx, page_img in enumerate(page_images):
        image_filename = f"{UPLOAD_DIR}/page_{idx}.png"
        page_img.save(image_filename, format='PNG')

        # Use OCR to extract text from image
        ocr_text = pytesseract.image_to_string(Image.open(image_filename))
        if ocr_text.strip():
            word_doc.add_paragraph("[Extracted from image]")
            word_doc.add_paragraph(ocr_text.strip())

        # Insert the page image into the Word document
        word_doc.add_picture(image_filename, width=Inches(5.5))

    # Save output Word file
    output_path = 'output.docx'
    word_doc.save(output_path)
    return output_path

if __name__ == '__main__':
    app.run(debug=True)
